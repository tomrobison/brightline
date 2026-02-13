"""
DOCX reading and writing utilities using python-docx.

Provides functions for reading resume content from DOCX files
and generating ATS-compatible formatted resumes.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any


def read_docx_paragraphs(file_path: str) -> List[Dict[str, Any]]:
    """
    Read all paragraphs from a DOCX file with formatting information.

    Args:
        file_path: Path to the DOCX file

    Returns:
        List of dictionaries containing paragraph text and metadata
    """
    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        para_dict = {
            'text': para.text,
            'style': para.style.name if para.style else 'Normal',
            'runs': []
        }

        # Extract run-level formatting
        for run in para.runs:
            para_dict['runs'].append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None
            })

        paragraphs.append(para_dict)

    return paragraphs


def create_ats_document() -> Document:
    """
    Create a new DOCX document with ATS-compatible formatting.

    Returns:
        Document object with proper ATS-safe settings
    """
    doc = Document()

    # Set narrow margins (0.5 inches all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    return doc


def add_contact_header(doc: Document, basics: Dict[str, Any]) -> None:
    """
    Add contact information header to resume.

    Args:
        doc: Document object
        basics: Dictionary with name, email, phone, location, etc.
    """
    # Name (larger, bold)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(basics.get('name', ''))
    name_run.bold = True
    name_run.font.size = Pt(16)

    # Contact info line
    contact_parts = []
    if basics.get('email'):
        contact_parts.append(basics['email'])
    if basics.get('phone'):
        contact_parts.append(basics['phone'])
    if basics.get('location', {}).get('city'):
        location = basics['location']
        loc_str = f"{location.get('city', '')}, {location.get('region', '')}"
        contact_parts.append(loc_str.strip(', '))

    if contact_parts:
        contact_para = doc.add_paragraph(' • '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)

    # LinkedIn/URL if available
    links = []
    if basics.get('url'):
        links.append(basics['url'])
    if basics.get('profiles'):
        for profile in basics['profiles']:
            if profile.get('network') == 'LinkedIn' and profile.get('url'):
                links.append(profile['url'])

    if links:
        link_para = doc.add_paragraph(' • '.join(links))
        link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        link_para.runs[0].font.size = Pt(9)

    # Add spacing after header
    doc.add_paragraph()


def add_section_header(doc: Document, title: str) -> None:
    """
    Add a section header (e.g., EXPERIENCE, EDUCATION).

    Args:
        doc: Document object
        title: Section title text
    """
    para = doc.add_paragraph()
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)

    # Add a subtle line under the header (using underline)
    run.underline = True


def add_experience_entry(doc: Document, experience: Dict[str, Any]) -> None:
    """
    Add a work experience entry to the resume.

    Args:
        doc: Document object
        experience: Dictionary with company, position, dates, bullets, etc.
    """
    # Company and Position line (bold)
    title_para = doc.add_paragraph()

    # Position in bold
    position_run = title_para.add_run(experience.get('position', ''))
    position_run.bold = True
    position_run.font.size = Pt(11)

    # Company (if different from position)
    if experience.get('company'):
        company_run = title_para.add_run(f" | {experience['company']}")
        company_run.font.size = Pt(11)

    # Location and Dates line
    info_parts = []
    if experience.get('location'):
        info_parts.append(experience['location'])

    # Format dates
    start = experience.get('startDate', '')
    end = experience.get('endDate', 'Present')
    if start:
        dates = f"{start} - {end}"
        info_parts.append(dates)

    if info_parts:
        info_para = doc.add_paragraph(' | '.join(info_parts))
        info_para.runs[0].font.size = Pt(10)
        info_para.runs[0].italic = True

    # Summary (if provided)
    if experience.get('summary'):
        summary_para = doc.add_paragraph(experience['summary'])
        summary_para.runs[0].font.size = Pt(10)

    # Bullet points
    if experience.get('bullets'):
        for bullet in experience['bullets']:
            bullet_para = doc.add_paragraph(bullet, style='List Bullet')
            bullet_para.runs[0].font.size = Pt(10)
            # Reduce spacing between bullets
            bullet_para.paragraph_format.space_before = Pt(0)
            bullet_para.paragraph_format.space_after = Pt(2)

    # Add spacing after experience
    doc.add_paragraph()


def add_skills_section(doc: Document, skills: List[Dict[str, Any]]) -> None:
    """
    Add skills section to resume.

    Args:
        doc: Document object
        skills: List of skill categories with keywords
    """
    for skill_category in skills:
        category = skill_category.get('category', '')
        keywords = skill_category.get('keywords', [])

        if not keywords:
            continue

        # Category name (bold)
        para = doc.add_paragraph()
        category_run = para.add_run(f"{category}: ")
        category_run.bold = True
        category_run.font.size = Pt(10)

        # Keywords (regular text)
        keywords_run = para.add_run(', '.join(keywords))
        keywords_run.font.size = Pt(10)


def add_education_entry(doc: Document, education: Dict[str, Any]) -> None:
    """
    Add education entry to resume.

    Args:
        doc: Document object
        education: Dictionary with institution, degree, dates, etc.
    """
    # Degree and Institution line
    title_para = doc.add_paragraph()

    # Degree in bold
    degree = f"{education.get('studyType', '')} in {education.get('area', '')}"
    degree_run = title_para.add_run(degree)
    degree_run.bold = True
    degree_run.font.size = Pt(11)

    # Institution
    if education.get('institution'):
        inst_run = title_para.add_run(f" | {education['institution']}")
        inst_run.font.size = Pt(11)

    # Dates and GPA
    info_parts = []
    if education.get('endDate'):
        info_parts.append(education['endDate'])
    if education.get('gpa'):
        info_parts.append(f"GPA: {education['gpa']}")

    if info_parts:
        info_para = doc.add_paragraph(' | '.join(info_parts))
        info_para.runs[0].font.size = Pt(10)
        info_para.runs[0].italic = True


def add_project_entry(doc: Document, project: Dict[str, Any]) -> None:
    """
    Add project entry to resume.

    Args:
        doc: Document object
        project: Dictionary with name, description, technologies, etc.
    """
    # Project name (bold)
    title_para = doc.add_paragraph()
    name_run = title_para.add_run(project.get('name', ''))
    name_run.bold = True
    name_run.font.size = Pt(11)

    # URL if available
    if project.get('url'):
        url_run = title_para.add_run(f" | {project['url']}")
        url_run.font.size = Pt(9)

    # Description
    if project.get('description'):
        desc_para = doc.add_paragraph(project['description'])
        desc_para.runs[0].font.size = Pt(10)

    # Technologies
    if project.get('technologies'):
        tech_para = doc.add_paragraph()
        tech_label = tech_para.add_run('Technologies: ')
        tech_label.italic = True
        tech_label.font.size = Pt(10)

        tech_text = tech_para.add_run(', '.join(project['technologies']))
        tech_text.font.size = Pt(10)


def estimate_page_count(doc: Document) -> float:
    """
    Estimate the number of pages in a document.

    This is a rough heuristic based on paragraph and character counts.

    Args:
        doc: Document object

    Returns:
        Estimated page count
    """
    total_chars = 0
    total_paragraphs = 0

    for para in doc.paragraphs:
        if para.text.strip():
            total_chars += len(para.text)
            total_paragraphs += 1

    # Rough heuristic: ~3000 characters per page, ~40 paragraphs per page
    char_estimate = total_chars / 3000
    para_estimate = total_paragraphs / 40

    # Average the two estimates
    return (char_estimate + para_estimate) / 2


def save_document(doc: Document, output_path: str) -> None:
    """
    Save document to file with proper error handling.

    Args:
        doc: Document object
        output_path: Path to save the document
    """
    try:
        doc.save(output_path)
        print(f"✓ Resume saved to: {output_path}")
    except Exception as e:
        print(f"✗ Error saving resume: {e}")
        raise
